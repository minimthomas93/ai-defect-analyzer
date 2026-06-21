# Main script for analyzing defects and generating bug reports
import os
import sys
from dotenv import load_dotenv
from google import genai
from google.genai import types
from docx import Document

def analyze_defects(input_bugfile):
    load_dotenv()
    #read the input bug written in the text file
    with open(input_bugfile, "r") as f:
        input_bug = f.read()
    print(input_bug)

    #create gemini client
    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

    #prompt for analyse the files to generate a structured report
    defect_analyse_prompt = f"""
        You are a Senior QA Engineer. Below are rough notes about a bug:

        BUG NOTES:
        {input_bug}

        Convert these into a professional bug report.
        Return EXACTLY in this format with ### before each section title:

        ### Title
        (content here)

        ### Summary
        (content here)

        ### Steps to Reproduce
        (content here)

        ### Expected Result
        (content here)

        ### Actual Result
        (content here)

        ### Environment
        (content here)

        ### Severity
        (content here)

        ### Priority
        (content here)

        ### Suggested Fix Area
        (content here)

        ### Missing Information
        (content here)

        ### Screenshots/Recordings URL
        (leave empty)

        ### Author
        (leave empty)

        - If the user mentions a severity, use it as a starting point but adjust if needed and explain why
    """

    api_response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=defect_analyse_prompt,
        config=types.GenerateContentConfig(
            temperature=0.1
    ))

    bug_content = api_response.text
    return bug_content

def generate_bugreport(bug_content,bugreport_file):
    doc = Document()        #create new word document
    doc.add_heading("BUG REPORT",level=0)       #add title

    sections = bug_content.split("### ")  #split the content into sections based on the ### delimiter

    for section in sections:
        if section:
            section = section.strip()
            lines = section.split("\n",1)
            headings = lines[0].strip()
            content = lines[1].strip() if len(lines) > 1 else ""

            doc.add_heading(headings, level=1)
            if content:
                doc.add_paragraph(content)

    doc.save(bugreport_file)    #save to disk
    print(f"BUG REPORT saved to {bugreport_file}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python defect_analyzer.py inputs/bug.txt")
        sys.exit(1)
    input_file = sys.argv[1]
    bug_content = analyze_defects(input_file)
    bugreport_file = os.path.join("outputs","bug_report.docx")
    generate_bugreport(bug_content,bugreport_file)
